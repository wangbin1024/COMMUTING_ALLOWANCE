from typing import Any, Generator, Optional, Sequence, Union, TypedDict, Literal, Annotated, Union, List, Dict
from pydantic import BaseModel, Field
from pyspark.sql import SparkSession
from databricks_langchain import ChatDatabricks, UCFunctionToolkit, VectorSearchRetrieverTool
from databricks.vector_search.client import VectorSearchClient
from databricks_langchain import DatabricksVectorSearch, ChatDatabricks
# import mlflow
from mlflow.langchain.chat_agent_langgraph import ChatAgentState, ChatAgentToolNode
from mlflow.pyfunc import ChatAgent
from mlflow.types.agent import ChatAgentChunk, ChatAgentMessage, ChatAgentResponse, ChatContext
from langchain.agents import create_tool_calling_agent, AgentExecutor
from langchain_core.language_models import LanguageModelLike
from langchain_core.prompts import PromptTemplate, ChatPromptTemplate, MessagesPlaceholder, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain_core.runnables import RunnableConfig, RunnableLambda
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage, SystemMessage
from langchain_core.tools import BaseTool, Tool
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser
from langchain_community.embeddings import HuggingFaceEmbeddings
from sentence_transformers import SentenceTransformer
from langgraph.graph import END, StateGraph 
from langgraph.types import Command, interrupt
from langgraph.graph.graph import CompiledGraph
from langgraph.graph.state import CompiledStateGraph
from langgraph.prebuilt.tool_node import ToolNode
from langgraph.checkpoint.memory import MemorySaver
from google.cloud import vision
from databricks import sql
from pypdf import PdfReader
from docx import Document
from datetime import date, datetime
import json
import requests
import pandas as pd
import re
import math
import os
import operator
import uuid

from langchain.callbacks import LangChainTracer

LANGSMITH_API_KEY = "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"
LANGSMITH_PROJECT = "agent-eval-project"

os.environ["LANGCHAIN_API_KEY"] = LANGSMITH_API_KEY
os.environ["LANGCHAIN_PROJECT"] = LANGSMITH_PROJECT

tracer = LangChainTracer()

############################################
# LLM endpoint
############################################
LLM_ENDPOINT_NAME_1 = "databricks-meta-llama-3-3-70b-instruct"
LLM_ENDPOINT_NAME_2 = "databricks-claude-3-7-sonnet"
llm = ChatDatabricks(endpoint=LLM_ENDPOINT_NAME_2, temperature=0.0)

############################################
# system prompt
############################################
system_prompt = """
あなたは業務支援のAIアシスタントです。
必要に応じて適切なツールを使って、ユーザーの質問や指示に対応してください。
"""

# 状態管理カスタマイズ　ChatAgentState拡張
class MyState(ChatAgentState):
  tools_ran: Annotated[set[str], operator.or_]

# COMMAND ----------

# utils関数
def convert_dates(obj):
  if isinstance(obj, (date, datetime)):
    return obj.isoformat()
  return obj

def safe_json(obj):
  if isinstance(obj, list):
    return [safe_json(item) for item in obj]
  elif isinstance(obj, dict):
    return {k: safe_json(v) for k, v in obj.items()}
  else:
    return convert_dates(obj)

def format_context(docs):
  chunk_template = (
    "{chunk_text}\n\n"
  )
  chunk_contents = [
    chunk_template.format(
      index=i + 1,
      chunk_text=d.page_content.strip()
    )
    for i, d in enumerate(docs)
  ]

  contents = "".join(chunk_contents)
  return contents

def vector_search(state: MyState) -> str:
    question = state["messages"][-2]["content"]
    print(f"\n📩 vector_search")
    # Connect to the Vector Search Index
    vs_client = VectorSearchClient(disable_notice=True)

    VECTOR_SEARCH_ENDPOINT = 'commuting_allowance_vector_search'
    VECTOR_SEARCH_INDEX = 'hhhd_demo_itec.commuting_allowance.commuting_allowance_index'

    # LangChain retrieverの作成
    embedding_model = HuggingFaceEmbeddings(
        model_name="/Workspace/Users/wang-b2@itec.hankyu-hanshin.co.jp/ruri-base-v2"
    )
    vector_search_as_retriever = DatabricksVectorSearch(
        endpoint=VECTOR_SEARCH_ENDPOINT,
        index_name=VECTOR_SEARCH_INDEX,
        embedding=embedding_model,
        text_column="chunked_text",
        columns=["chunk_id", "chunked_text"]
    ).as_retriever(search_kwargs={"k": 8, "score_threshold": 0.7})

    # result = vector_search_as_retriever.invoke(question)
    result = vector_search_as_retriever.get_relevant_documents(question)
    if len(result) > 0:
        context = format_context(result)
        tool_call = state["messages"][-1]["tool_calls"][0]
        arguments = json.loads(tool_call["function"]["arguments"])
        if "__arg1" in arguments:
            arguments["__arg1"] = context
            
        tool_call["function"]["arguments"] = json.dumps(arguments, ensure_ascii=False)
        return state
    return state

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, "rb") as f:
        reader = PdfReader(f)
        text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    text = ""

    # 通常のテキスト（段落）
    for para in doc.paragraphs:
        if para.text.strip():  # 空白行を削除
            text += para.text.strip() + "\n"

    # テーブルデータの取得
    for table in doc.tables:
        processed_cells = set()  # 既に処理したセルを保存

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if (cell._tc not in processed_cells) and cell_text:  # 結合セルを除外
                    row_data.append(cell_text)
                    processed_cells.add(cell._tc)
                else:
                    row_data.append("")

            text += " | ".join(row_data) + "\n"

    return text

def extract_text_from_xlsx(xlsx_path):
    df = pd.read_excel(xlsx_path, engine='openpyxl', sheet_name=None) 
    text = ""
    for sheet_name, sheet_df in df.items():
        text += f"### {sheet_name} シート ###\n"
        text += sheet_df.to_string(index=False, header=True) + "\n"
    return text

def extract_text_from_txt(txt_path):
    with open(txt_path, "r", encoding="shift_jis") as f:
        text = f.read()
    return text

def extract_text(file_path):
    text = os.path.splitext(file_path)[1].lower()

    if text == ".pdf":
        return extract_text_from_pdf(file_path)
    elif text == ".docx":
        return extract_text_from_docx(file_path)
    elif text in [".xls", ".xlsx"]:
        return extract_text_from_xlsx(file_path)
    elif text == ".txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"対応していないファイル形式です: {text}")

def create_llm_agent(model, tools, prompt):
  agent = create_tool_calling_agent(model, tools, prompt)
  return AgentExecutor(agent=agent, tools=tools, verbose=True)

# COMMAND ----------

############################################
# tools
############################################
def get_tax_adjustment_info(datas):
  try:
    # JSONからオブジェクトに変換
    data = json.loads(datas)
    user_id = data["ユーザーID"]
    insurance_classification = data["保険区分"]
    insurance_serial_number = data["保険内連番"]

    spark = SparkSession.builder.getOrCreate()

    query = f"""
    SELECT * FROM hhhd_demo_itec.tax_adjustment.`保険料明細情報` 
    WHERE `申請状況（保険料）` <> '完了' AND `ユーザーID` = '{user_id}' AND `保険区分` = '{insurance_classification}'  
    """

    if insurance_serial_number:
        query += f" AND `保険内連番` = '{insurance_serial_number}'"

    df = spark.sql(query)
    result = df.toPandas().to_dict(orient="records")

    print(f"申請情報result: {result}")

    return json.dumps({
      "申請情報": safe_json(result[0]) if result else {},
      "件数": len(result)
    }, ensure_ascii=False)

  except Exception as e:
    return json.dumps({
      "エラー": f"申請情報取得中にエラーが発生しました: {str(e)}"
    }, ensure_ascii=False)

def ocr_tool(image_name):
    try:
        # サービスアカウントキーを指定
        spark = SparkSession.builder.getOrCreate()
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "/Volumes/hhhd_demo_itec/tax_adjustment/keys/directed-reef-454701-d6-0f8e0d1de3c1.json"

        # Vision APIクライアントの作成
        client = vision.ImageAnnotatorClient()

        # OCR対象画像の読み込み
        image_path = f"/Volumes/hhhd_demo_itec/tax_adjustment/images/{image_name}"
        with open(image_path, "rb") as image_file:
            content = image_file.read()
            image = vision.Image(content=content)

        # OCR実行（日本語含む）
        response = client.text_detection(image=image)
        texts = response.text_annotations

        if response.error.message:
            raise Exception(f"APIエラー: {response.error.message}")

        if texts:
            vision_text = texts[0].description
            return json.dumps({
                "画像情報": {
                "ocr全文": vision_text.strip()
                }
            }, ensure_ascii=False)
        else:
            return json.dumps({
                "画像情報": {
                "ocr全文": "",
                "注意": "テキストが検出されませんでした"
                }
            }, ensure_ascii=False)

    except Exception as e:
        return json.dumps({
        "エラー": f"OCR処理中にエラーが発生しました: {str(e)}"
        }, ensure_ascii=False)

def check_tax_adjustment_consistency(datas):
    system_prompt = SystemMessagePromptTemplate.from_template("""
    あなたは年末調整の申請内容と、OCRで読み取った保険料控除証明書画像の情報を照合するAIアシスタントです。
    生命保険の場合は保険分類があります。同じ保険分類の情報を照合してください。
    保険料控除証明には11月分と12月分が分かれて記載される場合があります。この場合は申告額を合算してください。
    保険料控除証明書に複数申告年が書かれている場合は、同じ申告年の情報を照合してください。

    【タスク】
    以下の項目について、申請情報と画像情報を比較し、一致しているかどうかを判断してください。
    - 契約者名
    - 保険会社・事業所名
    - 保険名称
    - 保険期間
    - 生命保険分類（生命保険の場合のみ）
    - 保険種類
    - 申告額
    - 申告年

    【出力形式】
    以下のような読みやすい形式で、比較結果を記述してください：

    ---
    ## 契約者名：
    - 申請情報：XX 太郎
    - 画像情報：XX 太郎
    - 一致

    ## 保険名称：
    - 申請情報：XX保険
    - 画像情報：OO保険
    - 不一致

    ## 総合評価：
    ---
    """)
    user_prompt = HumanMessagePromptTemplate.from_template("{input}")
    prompt = ChatPromptTemplate.from_messages([
        system_prompt,
        user_prompt,
    ])
    model = ChatDatabricks(endpoint=LLM_ENDPOINT_NAME_2, temperature=0.0)
    extract_fields_chain = prompt | model | StrOutputParser()
    result = extract_fields_chain.invoke({"input": datas})
    return result

def call_ehr_api_remand(data):
    print("差戻処理開始")
    # response = requests.get("https://api.hhhd.jp/ehr/remand")

    # if response.error.message:
    #     raise Exception(f"APIエラー: {response.error.message}")
    # else:
    #     return response.text
    return "差戻しました"
    
def call_ehr_api_approval(data):
    print("承認処理開始")
    # response = requests.get("https://api.hhhd.jp/ehr/approval")

    # if response.error.message:
    #     raise Exception(f"APIエラー: {response.error.message}")
    # else:
    #     return response.text
    return "承認しました"

def search_company_regulations(question: str) -> str:
    return f"会社規定:\n\n{question}"

def get_tax_adjustment_history(name: str) -> str:
    spark = SparkSession.builder.getOrCreate()
    clean_name = name.replace(" ", "")
    if not clean_name:
        print("名前が指定されていないため、検索をスキップします。")
        return []

    query = f"""
    SELECT name, work_address, address, nearest_station, 
            route_1, distance_1, commuter_pass_1, 
            route_2, distance_2, commuter_pass_2
    FROM hhhd_demo_itec.allowance_payment_rules.commuting_allowance_history
    WHERE name = '{clean_name}'
    """

    try:
        df = spark.sql(query)
        text = json.dumps(df.collect(), ensure_ascii=False)
        return f"過去の申請履歴:\n\n{text}"
    except Exception as e:
        print(f"SQL実行中にエラーが発生しました: {e}")
        return "申請履歴が見つかりませんでした。"
 
def search_commute_routes(requests_data: str) -> str:
    try:
        data = json.loads(requests_data)
    except Exception as e:
        return json.dumps({"error": f"JSONのパースに失敗しました: {e}"})

    api_key = "test_nfY87YJYHMp"
    base_url = "https://api.ekispert.jp/v1/json/search/course/extreme"

    from_address = data.get("from_address")
    to_address = data.get("to_address")
    
    if not from_address or not to_address:
        return json.dumps({"error": "エラー: 出発地と到着地の住所は必須です。"})
    print(f"from_address: {from_address} to_address: {to_address}")
    radius = 1000
    result_data = []

    sort_types = {
        "time": "最短ルート",
        "teiki": "最安ルート",
        "transfer": "乗換少ないルート"
    }

    for sort_key, label in sort_types.items():
        params = {
            "key": api_key,
            "viaList": f"{from_address},{radius}:{to_address},{radius}",
            "searchType": "plain",
            "sort": sort_key,
            "answerCount": 10,
        }

        try:
            response = requests.get(base_url, params=params)
            if response.status_code != 200:
                result_data.append({"label": label, "error": f"APIエラー: {response.status_code}"})
                continue

            data = response.json()
            course = data.get("ResultSet", {}).get("Course")
            if isinstance(course, list):
                course = course[0]

            teiki = course.get("Teiki", {})
            displayRoute = teiki.get("DisplayRoute", {})
            route = course.get("Route", {})
            lines = route.get("Line", [])
            points = route.get("Point", [])

            timeOther = route.get("timeOther", "不明")
            timeOnBoard = route.get("timeOnBoard", "不明")
            timeWalk = route.get("timeWalk", "不明")
            distance = route.get("distance", "不明")
            transfer_count = route.get("transferCount", "不明")

            fare = "不明"
            teiki1 = teiki3 = teiki6 = "不明"
            for price in course.get("Price", []):
                kind = price.get("kind") or price.get("Kind")
                if kind == "FareSummary":
                    fare = price.get("Oneway", "不明")
                elif kind == "Teiki1Summary":
                    teiki1 = price.get("Oneway", "不明")
                elif kind == "Teiki3Summary":
                    teiki3 = price.get("Oneway", "不明")
                elif kind == "Teiki6Summary":
                    teiki6 = price.get("Oneway", "不明")

            line_info = []
            for i, line in enumerate(lines):
                dep_idx = i
                arr_idx = i + 1

                dep_name = "（出発地不明）"
                arr_name = "（到着地不明）"

                if dep_idx < len(points):
                    point_dep = points[dep_idx]
                    dep_name = (
                        point_dep.get("Station", {}).get("Name")
                        or point_dep.get("Name")
                        or dep_name
                    )

                if arr_idx < len(points):
                    point_arr = points[arr_idx]
                    arr_name = (
                        point_arr.get("Station", {}).get("Name")
                        or point_arr.get("Name")
                        or arr_name
                    )

                transport_name = line.get("Name") or line.get("TypicalName")

                if not transport_name or transport_name.strip() == "":
                    continue

                line_info.append({
                    "transport_name": transport_name,
                    "from": dep_name,
                    "to": arr_name
                })

            result_data.append({
                "label": label,
                "time_on_board": timeOnBoard,
                "time_other": timeOther,
                "time_walk": timeWalk,
                "total_time": int(timeOnBoard) + int(timeWalk) + int(timeOther),
                "distance": distance,
                "transfer_count": transfer_count,
                "fare": fare,
                "teiki_1": teiki1,
                "teiki_3": teiki3,
                "teiki_6": teiki6,
                "route": line_info
            })

        except Exception as e:
            result_data.append({"label": label, "error": f"例外エラー: {e}"})

    return json.dumps(result_data, ensure_ascii=False)

def extract_text_from_file(file_name: str) -> str:
    file_path = f"/Volumes/hhhd_demo_itec/commuting_allowance/inputs/{file_name}"

    try:
        extracted_text = extract_text(file_path)
        text = f"\n---\n\n### {os.path.basename(file_path)} の申請内容\n\n```\n{extracted_text}\n```\n"
    except Exception as e:
        print(f"エラー: {file_path} の処理に失敗しました: {e}")
        return ""

    if not text.strip():
        print("ファイルの読み取りに失敗した、もしくは内容が空でした。")
        return "ファイルの読み取りに失敗した、もしくは内容が空でした。"

    try:
        # questionから情報を抽出して、定型文に変換するChain　
        extract_prompt_template = PromptTemplate.from_template("""
            申請書データ:
            {input_text}                                                      

            上記の申請書データから、以下の情報を抽出してください：

            - 申請者名
            - 勤務先住所
            - 自宅住所
            - 最寄り駅
            - 利用交通機関と経路
            - 通勤距離
            - 定期代

            定期代が経路ごとに書かれている場合は、それぞれの定期代を足した総額を記載してください。
            出力は以下のMarkdown形式に従ってください：

            ---
            ##申請書内容##
            - **申請者名**:  
            - **勤務先住所**:  
            - **自宅住所**:  
            - **最寄り駅**: 
            - **定期代**: 
            - **利用交通機関と経路①**:  
            - **通勤距離①**:   
            ...（以下続く）
            ---
            """)
        model = ChatDatabricks(endpoint=LLM_ENDPOINT_NAME_2, temperature=0.0)
        extract_fields_chain = extract_prompt_template | model | StrOutputParser()

        llm_text = extract_fields_chain.invoke({"input_text": text})
    except Exception as e:
        print(f"LLMによる構造化抽出に失敗しました: {e}")
        return text.strip()

    return llm_text

def check_commuting_allowance_consistency(datas):
    system_prompt = SystemMessagePromptTemplate.from_template("""
    - あなたは会社の通勤手当の申請が適正かどうかをチェックするシステムです。  
    -【申請内容】を【会社規定】と照らし合わせて、申請が適切かどうかを判断してください。
    -【過去の申請履歴】がある場合は、その**申請履歴を表示すること**。履歴がない場合は、申請履歴のセクションは表示しないこと。
    -【最寄り駅と通勤ルートなどの検索結果】がある場合は、あわせて、判断すること。
    - 経路が複数に分かれている場合は、定期代は**合計金額*で判断すること。また、通勤距離は**合計距離**で判断すること。
    - 最終判断は、以下の3つから選び、**必ず理由と規定の引用を添えて出力すること。**
    - **問題なし**：すべての規定に適合している場合。
    - **問題あり**：一つでも規定に反している場合。
    - **要確認**：規定に該当がない、判断できない、または別途基準が必要な場合。
    ---

    ### **出力フォーマット**

    #### [申請者名] さん

    - **過去の申請履歴**:  
        - [申請履歴内容]  
        - or `なし`

    - **最寄り駅確認結果**:  
    - `一致`  
    - `不一致`　検索された最寄り駅：[検索された最寄り駅名]

    - **会社規定の引用と判断理由**:  
        - `[引用した会社規定]`  
        - `[規程に基づいた詳細な判断理由]`

    - **最終判断結果**:  
    - `問題なし`  
    - `問題あり`  
    - `要確認`
    """)
    user_prompt = HumanMessagePromptTemplate.from_template("{input}")
    prompt = ChatPromptTemplate.from_messages([
        system_prompt,
        user_prompt,
    ])
    model = ChatDatabricks(endpoint=LLM_ENDPOINT_NAME_2, temperature=0.0)
    extract_fields_chain = prompt | model | StrOutputParser()
    result = extract_fields_chain.invoke({"input": datas})
    return result

def call_commuting_allowance_api_remand(data):
    print("通勤手当申請差戻処理開始")
    return "通勤手当申請差戻しました"
    
def call_commuting_allowance_api_approval(data):
    print("通勤手当申請承認処理開始")
    return "通勤手当申請承認しました"

# args_schemas
class TaxAdjustmentArgs(BaseModel):
  adjustment_info: dict = Field(
    description="年末調整申請情報。契約者名、保険会社名、保険名称、年金年、保険証券番号などを含むJSONオブジェクト形式。"
  )
  image_info: dict = Field(
    description="OCRによって抽出された保険証明書画像の内容。全文または項目別の情報を含むJSON形式。"
  )

tools = [
    Tool(
        name="search_company_regulations",
        func=search_company_regulations,
        description="""
        会社規定を取得するツールです。
        「会社規定を参照」と言われた場合は、このツールを使用してください。
        """
    ),
    Tool(
        name="get_tax_adjustment_history",
        func=get_tax_adjustment_history,
        description="""
        通勤手当の過去の申請履歴を取得するツールです。
        名前を使って、過去の申請履歴を取得します。
        入力例: 「申請者の名前」
        """
    ),
    Tool(
        name="extract_text_from_file",
        func=extract_text_from_file,
        description="""
        申請書名を渡して、申請書内容を取得するツールです。
        入力例: "sample.pdf"
        """
    ),
    Tool(
        name="search_commute_routes",
        func=search_commute_routes,
        description= """
        JSON形式の文字列で渡された出発地と到着地の住所を使って、駅すぱあとAPIから最短・最速・最楽ルートを取得するツールです。
        注意: 出発地と到着地は必ず都道府県から始まる完全な日本語住所にしてください。そして、ビル名などが要りません。
        入力例: '{"from_address": "大阪府大阪市中央区1-1-1", "to_address": "大阪府大阪市福島区1-2-3"}'
        """
    ),
    Tool(
        name="call_commuting_allowance_api_remand",
        func=call_commuting_allowance_api_remand,
        description="""
        通勤手当申請の差戻APIを呼び出すツールです。
        申請内容を渡して、通勤手当申請の差戻APIを呼び出すと、差戻結果が返ります。
        """
    ),
    Tool(
        name="call_commuting_allowance_api_approval",
        func=call_commuting_allowance_api_approval,
        description="""
        通勤手当申請の承認APIを呼び出すツールです。
        申請内容を渡して、通勤手当申請の承認APIを呼び出すと、承認結果が返ります。
        """
    ),
    Tool(
        name="check_commuting_allowance_consistency",
        func=check_commuting_allowance_consistency,
        description="""
        通勤手当申請が適正かどうかをチェックするツールです。

        以下4つの引数を必ず指定してください：

        - 申請内容: 通勤手当申請書の申請内容
        - 会社規定: 通勤手当に関する会社規定
        - 過去の申請履歴: ユーザーが過去申請した通勤手当の申請履歴
        - 最寄り駅と通勤ルートなどの検索結果: 駅すぱあとAPIから取得した最短・最速・最楽ルートなどの検索結果

        例：
        {
        "申請内容": { ... },
        "会社規定": { ... },
        "過去の申請履歴": { ... },
        "最寄り駅と通勤ルートなどの検索結果": { ... },
        }
        """
    ),
    Tool(
        name="get_tax_adjustment_info",
        func=get_tax_adjustment_info,
        description="""
        "ユーザーID", "保険区分", "保険内連番"を使って、年末調整の申請情報を取得するツールです。
        必ずダブルクォートで囲まれた **有効なJSON文字列** を入力してください。
        入力例: "{\"ユーザーID\": 0000001, \"保険区分\": \"XX保険\", \"保険内連番\": 1}"
        """
    ),
    Tool(
        name="ocr_tool",
        func=ocr_tool,
        description="""
        画像名を渡して、画像情報を取得するツールです。
        入力例: "image.jpg"
        """
    ),
    Tool(
        name="call_ehr_api_remand",
        func=call_ehr_api_remand,
        description="""
        e-じんじの差戻APIを呼び出すツールです。
        申請内容を渡して、e-じんじの差戻APIを呼び出すと、差戻結果が返ります。
        """
    ),
    Tool(
        name="call_ehr_api_approval",
        func=call_ehr_api_approval,
        description="""
        e-じんじの承認APIを呼び出すツールです。
        申請内容を渡して、e-じんじの承認APIを呼び出すと、承認結果が返ります。
        """
    ),
    Tool(
        name="check_tax_adjustment_consistency",
        func=check_tax_adjustment_consistency,
        args_schemas=TaxAdjustmentArgs,
        description="""
        年末調整申請内容と保険証明書（画像）のOCR結果を比較し、一致しているか確認します。

        以下2つの引数を必ず指定してください：

        - 申請情報: 年末調整の申請情報（辞書形式）
        - 画像情報: OCRによって取得した画像情報（辞書形式）

        例：
        {
        "申請情報": { ... },
        "画像情報": { ... }
        }
        """
    )
]

# COMMAND ----------


############################################
# Human-in-the-loop
############################################
# 人間参与
def human_assistance(state: MyState):
    print("### 管理者の確認待ち ###")
    # "name": "check_tax_adjustment_consistency
    last_query = state["messages"][-2]["content"]
    print(f"\n📩 エージェントの判断結果:\n{last_query}")

    response_text = input("次のやるべきことを入力してください: ")

    # tools_ran を更新（該当ツールを削除）
    updated_tools_ran = set(state.get("tools_ran", set()))
    updated_tools_ran.discard("check_tax_adjustment_consistency")

    return {
        "messages": state["messages"] + [{
        "role": "user",
        "content": response_text,
        "id": str(uuid.uuid4())
        }],
        "tools_ran": updated_tools_ran
    }

# ツールノード呼び出しカスタマイズ
def custom_tool_node(state: dict) -> dict:
    # ツールノード実行（LangGraph内部からの呼び出し）
    tool_node = ChatAgentToolNode(tools)
    tool_result_state = tool_node.invoke(state, config=None)

    # 最新のメッセージ取得（AIMessage or dict）
    last_msg = state["messages"][-1]

    # tool_calls 取得（両対応：dict or AIMessage）
    if isinstance(last_msg, dict):
        tool_calls = last_msg.get("tool_calls", [])
    else:
        tool_calls = getattr(last_msg, "tool_calls", [])

    # すでに実行済みの tool のセット
    tools_ran = set(state.get("tools_ran", set()))

    # 実行されたツール名を tools_ran に追加
    for call in tool_calls:
        # call 自体も dict または pydantic model の可能性があるため両対応
        if isinstance(call, dict):
            tool_name = call.get("name")
        else:
            tool_name = getattr(call, "name", None)
        if tool_name:
            tools_ran.add(tool_name)

    # 更新後の状態に tools_ran を追加
    tool_result_state["tools_ran"] = tools_ran
    return tool_result_state

#####################
## Define agent logic
#####################
# LangGraphベースのエージェント構築
def create_tool_calling_agent(
    model: LanguageModelLike,
    tools: Union[ToolNode, Sequence[BaseTool]],
    system_prompt: Optional[str] = None,
) -> CompiledGraph:
    # モデルにツールをバインド
    model = model.bind_tools(tools)

    # Define the function that determines which node to go to
    # 次の状態を判断する関数
    def should_continue(state: MyState) -> str:
        last_msg = state["messages"][-1]
        tools_ran = state.get("tools_ran", set())
        if "tool_calls" in last_msg and last_msg["tool_calls"]:
            function_name = last_msg["tool_calls"][0]["function"]["name"]
            if function_name == "search_company_regulations":
                return "vector_search"
            return "tools"
        elif "check_tax_adjustment_consistency" in tools_ran:
            tools_ran.discard("check_tax_adjustment_consistency")
            return "human"
        elif "check_commuting_allowance_consistency" in tools_ran:
            tools_ran.discard("check_commuting_allowance_consistency")
            return "human"
        return "end"

    # 前処理（入力された MyState から "messages" を取り出し、必要ならシステムプロンプトを先頭に追加）
    if system_prompt:
        preprocessor = RunnableLambda(
            lambda state: [{"role": "system", "content": system_prompt}]
            + state["messages"]
        )
    else:
        preprocessor = RunnableLambda(lambda state: state["messages"])

    # パイプライン定義
    model_runnable = preprocessor | model
    
    def call_model(state: MyState, config: RunnableConfig):
        response = model_runnable.invoke(state, config)
        if response is None:
            return {"messages": []}

        return {"messages": [response]}

    # 新しい状態stateを定義
    workflow = StateGraph(MyState)
    # nodeを追加
    workflow.add_node("agent", RunnableLambda(call_model))
    workflow.add_node("tools", RunnableLambda(custom_tool_node))
    workflow.add_node("human", RunnableLambda(human_assistance))
    workflow.add_node("vector_search", RunnableLambda(vector_search))

    ###########################
    ## 状態遷移ルールを定義
    ## agent → "tool_calls"あり → tools
    ## tools → 呼び出し後 → agent に戻る（再応答）
    ###########################
    # 入口
    workflow.set_entry_point("agent")
    # 条件edge追加
    workflow.add_conditional_edges(
        "agent",
        should_continue,
        {
            "tools": "tools",
            "human": "human",
            "vector_search": "vector_search",
            "end": END,
        },
    )
    # 一般edge追加
    workflow.add_edge("tools", "agent")
    workflow.add_edge("human", "agent")
    workflow.add_edge("vector_search", "tools")
    # メモリを追加
    memory = MemorySaver()
    # コンパイル
    return workflow.compile(checkpointer=memory)

# 出力結果の形を整形する
class LangGraphChatAgent(ChatAgent):
    def __init__(self, agent: CompiledStateGraph):
        self.agent = agent

    def predict(
        self,
        messages: list[ChatAgentMessage],
        context: Optional[ChatContext] = None,
        custom_inputs: Optional[dict[str, Any]] = None,
    ) -> ChatAgentResponse:
        request = {"messages": self._convert_messages_to_dict(messages)}
        config = {"configurable": {"thread_id": context.thread_id if context else "session_001"}, "callbacks": [tracer]}
        messages = []
        for event in self.agent.stream(request, config=config, stream_mode="updates"):
            for node_data in event.values():
                messages.extend(
                    ChatAgentMessage(**msg) for msg in node_data.get("messages", [])
                )
        return ChatAgentResponse(messages=messages)

    def predict_stream(
        self,
        messages: list[ChatAgentMessage],
        context: Optional[ChatContext] = None,
        custom_inputs: Optional[dict[str, Any]] = None,
    ) -> Generator[ChatAgentChunk, None, None]:
        request = {"messages": self._convert_messages_to_dict(messages)}
        config = {"configurable": {"thread_id": context.thread_id if context else "session_001"}, "callbacks": [tracer]}
        for event in self.agent.stream(request, config=config, stream_mode="updates"):
            for node_data in event.values():
                yield from (
                    ChatAgentChunk(**{"delta": msg}) for msg in node_data["messages"]
                )


# Create the agent object, and specify it as the agent object to use when
# loading the agent back for inference via mlflow.models.set_model()
# mlflow.langchain.autolog()
agent = create_tool_calling_agent(llm, tools, system_prompt)
AGENT = LangGraphChatAgent(agent)
# mlflow.models.set_model(AGENT)
